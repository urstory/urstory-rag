"""파일 변환 서비스: PDF, DOCX, TXT, Markdown → 텍스트 추출."""
import os
from dataclasses import dataclass, field
from pathlib import Path

from haystack.components.converters import (
    MarkdownToDocument,
    PyPDFToDocument,
    TextFileToDocument,
)
from haystack.dataclasses import ByteStream


@dataclass
class ConversionResult:
    content: str
    meta: dict = field(default_factory=dict)


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

# 확장자 → Haystack 컨버터 매핑
_CONVERTERS = {
    "pdf": lambda: PyPDFToDocument(),
    "txt": lambda: TextFileToDocument(),
    "md": lambda: MarkdownToDocument(),
}


class DocumentConverter:
    """지원 형식(PDF, DOCX, TXT, MD)의 파일을 텍스트로 변환."""

    def __init__(
        self,
        pdf_parser: str = "docling",
        ocr_enabled: bool = False,
        ocr_languages: list[str] | None = None,
        table_extraction_enabled: bool = True,
    ):
        self.pdf_parser = pdf_parser
        self.ocr_enabled = ocr_enabled
        self.ocr_languages = ocr_languages or ["ko", "en"]
        self.table_extraction_enabled = table_extraction_enabled

    def detect_file_type(self, filename: str) -> str:
        ext = Path(filename).suffix.lower().lstrip(".")
        return ext

    async def convert(self, file_path: str) -> ConversionResult:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")

        file_type = ext.lstrip(".")
        file_size = path.stat().st_size

        # Docling으로 PDF/DOCX 변환
        if file_type in ("pdf", "docx") and self.pdf_parser == "docling":
            return await self._convert_with_docling(path, file_type, file_size)

        content = await self._extract_content(path, file_type)

        return ConversionResult(
            content=content,
            meta={
                "filename": path.name,
                "file_type": file_type,
                "file_size": file_size,
            },
        )

    async def _convert_with_docling(
        self, path: Path, file_type: str, file_size: int
    ) -> ConversionResult:
        """Docling을 사용한 PDF/DOCX 변환."""
        from app.services.document.docling_converter import DoclingPDFConverter

        docling = DoclingPDFConverter(
            ocr_enabled=self.ocr_enabled,
            ocr_languages=self.ocr_languages,
            table_extraction_enabled=self.table_extraction_enabled,
        )
        return await docling.convert(str(path))

    async def _extract_content(self, path: Path, file_type: str) -> str:
        if file_type == "docx":
            return await self._convert_docx(path)

        converter_factory = _CONVERTERS.get(file_type)
        if not converter_factory:
            raise ValueError(f"지원하지 않는 파일 형식입니다: .{file_type}")

        converter = converter_factory()
        source = ByteStream.from_file_path(path)
        source.meta["file_path"] = str(path)
        result = converter.run(sources=[source])

        documents = result.get("documents", [])
        if not documents:
            return ""

        return "\n".join(doc.content for doc in documents if doc.content)

    async def _convert_docx(self, path: Path) -> str:
        """DOCX 변환 (python-docx 직접 사용)."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx 패키지가 필요합니다: pip install python-docx")

        doc = DocxDocument(str(path))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
